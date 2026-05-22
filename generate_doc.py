"""
generate_doc.py
---------------
Genere explication_detection_aberrantes.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code):
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def shading_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

BLUE  = (0x1e, 0x40, 0xaf)
BLUE2 = (0x16, 0x5a, 0xa0)

# ────────────────────────────────────────────────────────────────────────────
# PAGE DE GARDE
# ────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INNOV | CCNB")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Stage de fin d'etudes - Systeme d'alertes IoT")
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_heading("Detection des valeurs aberrantes", 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.color.rgb = RGBColor(*BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NumPy (Z-score, IQR) + scikit-learn (Isolation Forest)")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Export complet/nettoye et comparaison visuelle des donnees IoT")
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Date : " + datetime.date.today().strftime("%d/%m/%Y"))
run.font.size = Pt(11)

doc.add_page_break()

# ────────────────────────────────────────────────────────────────────────────
# SECTION 1 - CONTEXTE
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Contexte et objectif", 1, BLUE)

doc.add_paragraph(
    "Dans le cadre du systeme d'alertes IoT developpe lors de ce stage, les capteurs "
    "(temperature, turbidite, pH) transmettent des mesures en temps reel. "
    "Certaines de ces mesures peuvent etre aberrantes — dues a un bruit electronique, "
    "une panne passagere ou une transmission corrompue. "
    "Mon encadreur a demande d'utiliser NumPy et scikit-learn pour implementer "
    "des methodes statistiques et d'apprentissage automatique de detection."
)

doc.add_paragraph("L'objectif etait double :")
add_bullet(doc, "Detecter automatiquement les valeurs hors-norme en temps reel lors de la reception des donnees capteur.", "1.")
add_bullet(doc, "Permettre l'export des donnees en deux versions : completes et nettoyees (CSV, Excel, PDF).", "2.")

# ────────────────────────────────────────────────────────────────────────────
# SECTION 2 - NUMPY
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Methodes statistiques avec NumPy", 1, BLUE)

doc.add_paragraph(
    "NumPy (Numerical Python) est une bibliotheque de calcul scientifique. "
    "Elle permet d'effectuer des operations mathematiques vectorisees sur des tableaux de donnees. "
    "Dans ce projet, NumPy est utilise pour calculer en temps reel les statistiques "
    "d'une fenetre glissante de 30 lectures par capteur."
)

add_heading(doc, "2.1 Z-score", 2, BLUE2)

doc.add_paragraph(
    "Le Z-score mesure l'ecart d'une valeur par rapport a la moyenne de son groupe, "
    "exprime en nombre d'ecarts-types. Plus le Z-score est eleve en valeur absolue, "
    "plus la valeur est eloignee de la normale."
)

add_code_block(doc, """\
import numpy as np
from collections import deque

class ZScoreDetector:
    def __init__(self, window=30, warn_th=2.0, crit_th=3.0):
        self._windows = {}
        self.warn_th  = warn_th      # seuil WARNING  : |z| > 2
        self.crit_th  = crit_th      # seuil CRITICAL : |z| > 3

    def analyze(self, sensor, value):
        win = self._windows.setdefault(sensor, deque(maxlen=30))
        arr  = np.array(win)         # conversion en tableau NumPy
        mean = arr.mean()            # moyenne de la fenetre
        std  = arr.std()             # ecart-type
        win.append(value)
        if std == 0: return "NORMAL"
        z = abs((value - mean) / std)
        if z > self.crit_th:  return "CRITICAL"
        if z > self.warn_th:  return "WARNING"
        return "NORMAL"\
""")

doc.add_paragraph("")
add_bullet(doc, "np.array() convertit la fenetre en tableau vectorise traite en C.", "arr.mean() :")
add_bullet(doc, "arr.std() calcule l'ecart-type de dispersion des valeurs.", "arr.std() :")
add_bullet(doc, "Si |z| > 3 : CRITICAL. Si |z| > 2 : WARNING. Sinon : NORMAL.", "z-score :")

add_heading(doc, "2.2 IQR (ecart interquartile)", 2, BLUE2)

doc.add_paragraph(
    "L'IQR est plus robuste que le Z-score car il est insensible aux extremes "
    "deja presents dans la fenetre. Il definit des bornes basees sur les quartiles Q1 et Q3."
)

add_code_block(doc, """\
class IQRDetector:
    def analyze(self, sensor, value):
        arr = np.array(window)
        q1  = np.percentile(arr, 25)   # 1er quartile
        q3  = np.percentile(arr, 75)   # 3e quartile
        iqr = q3 - q1
        lo_warn = q1 - 1.5 * iqr      # borne basse WARNING
        hi_warn = q3 + 1.5 * iqr      # borne haute WARNING
        lo_crit = q1 - 3.0 * iqr      # borne basse CRITICAL
        hi_crit = q3 + 3.0 * iqr      # borne haute CRITICAL
        if value < lo_crit or value > hi_crit: return "CRITICAL"
        if value < lo_warn or value > hi_warn: return "WARNING"
        return "NORMAL"\
""")

doc.add_paragraph("")
add_bullet(doc, "np.percentile(arr, 25) : 1er quartile (25% des valeurs en dessous).", "Q1 :")
add_bullet(doc, "np.percentile(arr, 75) : 3e quartile.", "Q3 :")
add_bullet(doc, "IQR = Q3 - Q1. Bornes WARNING : x1.5. Bornes CRITICAL : x3.0.", "IQR :")

add_heading(doc, "2.3 Pourquoi NumPy plutot que Python pur ?", 2, BLUE2)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Critere"
hdr[1].text = "Python pur"
hdr[2].text = "NumPy"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
    shading_cell(cell, "DBEAFE")

for r in [
    ("Vitesse de calcul",    "Boucle Python (lente)",           "Vectorise en C (tres rapide)"),
    ("Moyenne / Ecart-type", "sum()/len(), formule manuelle",   ".mean(), .std() integrees"),
    ("Percentiles (IQR)",    "Tri + indexation manuelle",       "np.percentile() en 1 appel"),
    ("Lisibilite du code",   "Verbeux, risque d'erreur",        "Concis et expressif"),
    ("Fiabilite numerique",  "Accumulation d'erreurs float",    "Algorithmes stables"),
]:
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = r
    shading_cell(row[2], "EFF6FF")

doc.add_paragraph("")

# ────────────────────────────────────────────────────────────────────────────
# SECTION 3 - ISOLATION FOREST (SCIKIT-LEARN)
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Apprentissage automatique avec scikit-learn", 1, BLUE)

doc.add_paragraph(
    "scikit-learn est la bibliotheque de reference pour le Machine Learning en Python. "
    "Elle fournit des algorithmes pre-implementes, optimises et documentes. "
    "J'ai utilise l'algorithme Isolation Forest pour completer la detection statistique "
    "par une approche d'apprentissage non-supervise."
)

add_heading(doc, "3.1 Principe de l'Isolation Forest", 2, BLUE2)

doc.add_paragraph(
    "L'Isolation Forest fonctionne en construisant 100 arbres de decision aleatoires. "
    "A chaque arbre, l'algorithme choisit aleatoirement un attribut et une valeur de coupure "
    "pour isoler les points de donnees :"
)
add_bullet(doc, "Une valeur NORMALE necessite de nombreuses coupures pour etre isolee (chemin long).")
add_bullet(doc, "Une valeur ABERRANTE est isolee rapidement car elle est rare et extreme (chemin court).")
add_bullet(doc, "Le score est la longueur moyenne du chemin : negatif = anomalie, positif = normal.")

doc.add_paragraph(
    "Contrairement a Z-score et IQR qui supposent une distribution gaussienne ou symetrique, "
    "l'Isolation Forest ne fait aucune hypothese sur la forme de la distribution. "
    "Il est particulierement efficace pour detecter des anomalies multidimensionnelles "
    "et les comportements inhabituels difficiles a capter par des seuils fixes."
)

add_heading(doc, "3.2 Implementation dans le projet", 2, BLUE2)

add_code_block(doc, """\
from sklearn.ensemble import IsolationForest
import numpy as np
from collections import deque

class IsolationForestDetector:
    def __init__(self, window_size=100, min_samples=20,
                 contamination=0.05, retrain_every=20):
        self.contamination = contamination   # 5% d'anomalies attendues
        self._windows = {}                   # fenetre par capteur
        self._models  = {}                   # modele entraine par capteur
        self._counts  = {}                   # compteur pour reentrain.

    def _train(self, sensor):
        arr = np.array(self._windows[sensor]).reshape(-1, 1)
        model = IsolationForest(
            n_estimators  = 100,     # 100 arbres d'isolation
            contamination = self.contamination,
            random_state  = 42,      # reproductible
        )
        model.fit(arr)               # apprentissage non-supervise
        self._models[sensor] = model

    def analyze(self, sensor, value):
        self._windows.setdefault(sensor, deque(maxlen=100)).append(value)
        self._counts[sensor] = self._counts.get(sensor, 0) + 1

        if len(self._windows[sensor]) < 20:
            return "NORMAL"          # pas encore assez de donnees

        # Reentraine tous les 20 nouveaux points
        if sensor not in self._models or self._counts[sensor] % 20 == 0:
            self._train(sensor)

        score = self._models[sensor].decision_function([[value]])[0]
        pred  = self._models[sensor].predict([[value]])[0]  # -1 ou 1

        if pred == -1:
            if score < -0.1: return "CRITICAL"
            return "WARNING"
        return "NORMAL"\
""")

doc.add_paragraph("")
add_bullet(doc, "model.fit(arr) : entraine le modele sur la fenetre de 100 lectures (non-supervise).", "fit() :")
add_bullet(doc, "decision_function() : retourne le score d'anomalie (negatif = suspect).", "score :")
add_bullet(doc, "predict() : retourne -1 (anomalie) ou +1 (normal) selon le seuil contamination.", "predict() :")
add_bullet(doc, "score < -0.1 -> CRITICAL (tres isole). score < 0 -> WARNING. score > 0 -> NORMAL.", "seuils :")

add_heading(doc, "3.3 Comparaison des trois methodes", 2, BLUE2)

table3 = doc.add_table(rows=1, cols=4)
table3.style = "Table Grid"
hdr3 = table3.rows[0].cells
for i, txt in enumerate(["Critere", "Z-score", "IQR", "Isolation Forest"]):
    hdr3[i].text = txt
    hdr3[i].paragraphs[0].runs[0].bold = True
    shading_cell(hdr3[i], "1E3A8A")
    hdr3[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for r in [
    ("Bibliotheque",         "NumPy",            "NumPy",            "scikit-learn"),
    ("Type",                 "Statistique",      "Statistique",      "Machine Learning"),
    ("Hypothese",            "Distribution norm.","Symetrie IQR",    "Aucune"),
    ("Sensible aux extremes","Oui",              "Non (robuste)",    "Non (robuste)"),
    ("Apprentissage",        "Non",              "Non",              "Oui (non-supervise)"),
    ("Fenetre",              "30 lectures",      "30 lectures",      "100 lectures"),
    ("CRITICAL si",          "|z| > 3",          "Hors 3*IQR",       "score < -0.1"),
]:
    row = table3.add_row().cells
    row[0].text, row[1].text, row[2].text, row[3].text = r
    shading_cell(row[3], "EFF6FF")

doc.add_paragraph("")

# ────────────────────────────────────────────────────────────────────────────
# SECTION 4 - MOTEUR COMBINE
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. Moteur de detection combine", 1, BLUE)

doc.add_paragraph(
    "Le moteur de detection (AnomalyDetectionEngine) orchestre les 4 detecteurs "
    "en parallele sur chaque lecture et retourne le niveau le plus severe."
)

add_code_block(doc, """\
class AnomalyDetectionEngine:
    def __init__(self, window_size=30):
        self.rules   = RuleBasedDetector()          # seuils metier fixes
        self.zscore  = ZScoreDetector()             # NumPy Z-score
        self.iqr     = IQRDetector()                # NumPy IQR
        self.iforest = IsolationForestDetector()    # scikit-learn

    def analyze(self, reading):
        candidates = [
            self.rules.analyze(reading),
            self.zscore.analyze(reading),
            self.iqr.analyze(reading),
            self.iforest.analyze(reading),
        ]
        # Strategie worst-case : retenir le niveau le plus severe
        return max(candidates, key=lambda r: SEVERITY[r.level])\
""")

doc.add_paragraph("")
doc.add_paragraph(
    "Exemple : si les regles metier retournent WARNING et l'Isolation Forest retourne CRITICAL, "
    "le systeme enregistre CRITICAL. Ce mecanisme garantit qu'aucune anomalie "
    "ne passe entre les mailles des differentes methodes."
)

# ────────────────────────────────────────────────────────────────────────────
# SECTION 5 - EXPORT
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Export des donnees : complet vs nettoye", 1, BLUE)

doc.add_paragraph(
    "Chaque lecture est stockee avec son niveau detecte (NORMAL, WARNING, CRITICAL). "
    "L'utilisateur peut exporter en deux versions depuis la page /logs :"
)
add_bullet(doc, "Donnees completes : toutes les lectures, y compris aberrantes.")
add_bullet(doc, "Donnees nettoyees : uniquement les lectures de niveau NORMAL.")

doc.add_paragraph("Filtrage cote serveur via SQLAlchemy :")

add_code_block(doc, """\
def get_logs(self, ..., exclude_outliers=False):
    q = self.db.query(SensorLog)
    if exclude_outliers:
        q = q.filter(SensorLog.level == "NORMAL")
    return q.order_by(desc(SensorLog.created_at)).limit(limit).all()\
""")

doc.add_paragraph("")
doc.add_paragraph("Chaque fichier telecharge contient un message informatif :")
add_bullet(doc, "CSV  : lignes de commentaire # avec le nombre de lignes retirees.")
add_bullet(doc, "Excel: derniere ligne recapitulative avec le comptage.")
add_bullet(doc, "PDF  : sous-titre avec le bilan de nettoyage.")

# ────────────────────────────────────────────────────────────────────────────
# SECTION 6 - PAGE COMPARAISON
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Page de comparaison visuelle (/logs/compare)", 1, BLUE)

doc.add_paragraph(
    "Une page dediee permet de visualiser cote a cote les statistiques "
    "des donnees completes et nettoyees pour chaque capteur :"
)
add_bullet(doc, "Tableau : total, valeurs aberrantes, taux d'aberration par capteur.")
add_bullet(doc, "Graphique Chart.js : barres comparant complet vs nettoye.")
add_bullet(doc, "KPIs globaux : total, aberrantes, taux global.")
add_bullet(doc, "Interroge directement la base de donnees — instantane et toujours a jour.")

# ────────────────────────────────────────────────────────────────────────────
# SECTION 7 - COMPARAISON DEUX FICHIERS
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. Comparaison de deux fichiers exportes", 1, BLUE)

doc.add_paragraph(
    "L'utilisateur peut uploader deux fichiers CSV ou Excel pour comparer "
    "deux periodes differentes ou avant/apres une intervention."
)
add_bullet(doc, "Endpoint POST /api/logs/compare-two : accepte deux UploadFile.")
add_bullet(doc, "Calcule par capteur : total, aberrantes, taux, et delta entre fichiers.")
add_bullet(doc, "Delta negatif (moins d'aberrantes) = vert. Delta positif = rouge.")
add_bullet(doc, "Export PDF de la comparaison en un clic.")

# ────────────────────────────────────────────────────────────────────────────
# SECTION 8 - TABLEAU RECAPITULATIF
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "8. Recapitulatif des fonctionnalites implementees", 1, BLUE)

table2 = doc.add_table(rows=1, cols=3)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
for i, txt in enumerate(["Fonctionnalite", "Technologies", "Fichier principal"]):
    hdr2[i].text = txt
    hdr2[i].paragraphs[0].runs[0].bold = True
    shading_cell(hdr2[i], "1E3A8A")
    hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for r in [
    ("Detection Z-score",        "NumPy (mean, std)",                 "src/detection/statistical.py"),
    ("Detection IQR",            "NumPy (percentile)",                "src/detection/statistical.py"),
    ("Detection Isolation Forest","scikit-learn (IsolationForest)",   "src/detection/statistical.py"),
    ("Moteur combine (4 detect.)","Python worst-case",                "src/detection/engine.py"),
    ("Stockage niveau detecte",  "SQLAlchemy + SQLite",               "src/alerts/models.py"),
    ("Export CSV nettoye",       "FastAPI StreamingResponse",         "src/api/routers.py"),
    ("Export Excel nettoye",     "ExcelJS (navigateur)",              "logs.html"),
    ("Export PDF nettoye",       "jsPDF + autoTable",                 "logs.html"),
    ("Page comparaison BD",      "FastAPI + Chart.js",                "dashboard.py / logs_compare.html"),
    ("Comparaison 2 fichiers",   "FastAPI + openpyxl + pandas",       "routers.py / logs_compare.html"),
]:
    row = table2.add_row().cells
    row[0].text, row[1].text, row[2].text = r

doc.add_paragraph("")

# ────────────────────────────────────────────────────────────────────────────
# CONCLUSION
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Conclusion", 1, BLUE)

doc.add_paragraph(
    "L'utilisation combinee de NumPy et scikit-learn repond directement a la demande "
    "de l'encadreur. Les methodes Z-score et IQR (NumPy) couvrent les anomalies "
    "statistiques classiques, tandis que l'Isolation Forest (scikit-learn) apporte "
    "une couche d'apprentissage automatique non-supervise capable de detecter des "
    "comportements anormaux sans hypothese sur la distribution des donnees."
)
doc.add_paragraph("Le systeme complet permet :")
add_bullet(doc, "Une detection en temps reel avec 4 methodes complementaires.")
add_bullet(doc, "Un export des donnees dans 3 formats avec bilan de nettoyage.")
add_bullet(doc, "Une comparaison visuelle instantanee via le dashboard.")
add_bullet(doc, "Une comparaison de deux fichiers pour evaluer l'evolution dans le temps.")

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run(
    "Ces fonctionnalites constituent une reponse complete et professionnelle "
    "a la demande de traitement des donnees aberrantes dans un systeme IoT industriel."
)
run.italic = True
run.font.color.rgb = RGBColor(*BLUE)

# ── Sauvegarde ───────────────────────────────────────────────────────────────
output_path = "explication_detection_aberrantes.docx"
doc.save(output_path)
print(f"Document genere : {output_path}")
