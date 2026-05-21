"""
add_journal_entry.py
--------------------
Ajoute l'entree du jour dans journal_de_bord.docx (non-interactif).
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

JOURNAL_PATH = Path(__file__).parent / "journal_de_bord.docx"

NIVEAU_COLORS = {
    "NORMAL":    RGBColor(0x27, 0xAE, 0x60),
    "ATTENTION": RGBColor(0xF3, 0x9C, 0x12),
    "BLOQUE":    RGBColor(0xE7, 0x4C, 0x3C),
}

entree = {
    "taches": [
        "Implementation de la detection des valeurs aberrantes avec NumPy (Z-score et IQR)",
        "Ajout du parametre exclude_outliers dans get_logs() et count_logs() (SQLAlchemy)",
        "Creation de l'endpoint GET /api/logs/export (CSV, Excel, PDF) en version complete et nettoyee",
        "Ajout d'un message dans chaque fichier exporte indiquant le nombre de lignes aberrantes retirees",
        "Correction du bouton 'Historique des notifications' (largeur complete, bien visible)",
        "Creation de la page /logs/compare : comparaison visuelle donnees completes vs nettoyees (Chart.js)",
        "Ajout de l'upload de deux fichiers CSV/Excel pour comparaison cote a cote avec delta colore",
        "Refonte complete du design CSS : Inter, glassmorphism, gradients, ombres modernes",
        "Generation du document Word explicatif explication_detection_aberrantes.docx",
    ],
    "difficultes": [
        "Export : le parametre limit=5000 depassait la validation le=500 de l'API -> corrige en le=10000",
        "URLSearchParams : un meme objet partage donnait des valeurs dupliquees -> utilise 3 objets independants",
        "Generation Word : heredoc bash echouait avec les apostrophes -> ecrit le script dans un .py d'abord",
        "Nouveau template /logs/compare : 404 car serveur non rechargee -> redemarrage manuel necessaire",
    ],
    "prochaines": [
        "Tester les exports CSV/Excel/PDF sur les donnees reelles du serveur",
        "Presenter la page /logs/compare et la comparaison deux fichiers a l'encadreur",
        "Presenter le document Word explicatif au professeur",
    ],
    "notes": (
        "L'encadreur avait demande d'utiliser NumPy et scikit-learn pour la detection des aberrantes. "
        "NumPy a ete utilise pour Z-score (mean, std) et IQR (percentile) sur fenetre glissante de 30 lectures. "
        "Le moteur combine regles metier + Z-score + IQR retenant le niveau le plus severe. "
        "Le document Word explication_detection_aberrantes.docx est disponible a la racine du projet."
    ),
    "niveau": "NORMAL",
}


def ajouter_entree(doc, entree):
    today = date.today().strftime("%A %d %B %Y").capitalize()
    doc.add_paragraph("-" * 60)
    h = doc.add_heading(f"Journee du {today}", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x9C)

    doc.add_heading("Taches accomplies", level=3)
    for tache in entree["taches"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(tache)

    if entree.get("difficultes"):
        doc.add_heading("Difficultes rencontrees", level=3)
        for d in entree["difficultes"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(d)
            run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

    if entree.get("prochaines"):
        doc.add_heading("Prochaines etapes", level=3)
        for etape in entree["prochaines"]:
            doc.add_paragraph(etape, style="List Bullet")

    if entree.get("notes"):
        doc.add_heading("Notes", level=3)
        doc.add_paragraph(entree["notes"])

    niveau = entree.get("niveau", "NORMAL").upper()
    p = doc.add_paragraph()
    run = p.add_run(f"Niveau d'avancement : {niveau}")
    run.bold = True
    run.font.color.rgb = NIVEAU_COLORS.get(niveau, RGBColor(0, 0, 0))
    doc.add_paragraph("")


if not JOURNAL_PATH.exists():
    print(f"Fichier introuvable : {JOURNAL_PATH}")
else:
    doc = Document(JOURNAL_PATH)
    ajouter_entree(doc, entree)
    doc.save(JOURNAL_PATH)
    print(f"Journal mis a jour : {JOURNAL_PATH}")
