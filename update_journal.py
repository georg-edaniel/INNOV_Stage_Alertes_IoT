"""
update_journal.py
-----------------
Script de mise à jour quotidienne :
  1. Ajoute une entrée dans journal_de_bord.docx
  2. Commit + push automatique sur GitHub

Usage:
    python update_journal.py
"""

import subprocess
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installation de python-docx...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

JOURNAL_PATH = Path(__file__).parent / "journal_de_bord.docx"

NIVEAU_COLORS = {
    "NORMAL":    RGBColor(0x27, 0xAE, 0x60),   # vert
    "ATTENTION": RGBColor(0xF3, 0x9C, 0x12),   # orange
    "BLOQUÉ":    RGBColor(0xE7, 0x4C, 0x3C),   # rouge
}


def init_journal() -> Document:
    """Crée le fichier Word s'il n'existe pas encore."""
    doc = Document()

    # Titre principal
    titre = doc.add_heading("Journal de Bord — Stage INNOV", level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Infos
    doc.add_paragraph("Stagiaire : Philip Daniel George Moumie Gouet")
    doc.add_paragraph("Projet : Système d'Alertes Intelligentes — Données IoT Simulées")
    doc.add_paragraph("Période : 1er mai 2026 → 16 juin 2026")
    doc.add_paragraph("Lieu : INNOV / CCNB — Fabrication Avancée, Bathurst (N.-B.)")
    doc.add_paragraph("")

    doc.save(JOURNAL_PATH)
    print(f"Journal créé : {JOURNAL_PATH}")
    return doc


def ajouter_entree(doc: Document, entree: dict) -> None:
    """Ajoute une entrée quotidienne dans le document Word."""
    today = date.today().strftime("%A %d %B %Y").capitalize()

    # Séparateur + date
    doc.add_paragraph("─" * 60)
    h = doc.add_heading(f"Journée du {today}", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x9C)

    # Tâches accomplies
    doc.add_heading("✅ Tâches accomplies", level=3)
    for tache in entree["taches"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(tache)

    # Difficultés
    if entree.get("difficultes"):
        doc.add_heading("⚠️ Difficultés rencontrées", level=3)
        for d in entree["difficultes"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(d)
            run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

    # Prochaines étapes
    if entree.get("prochaines"):
        doc.add_heading("🎯 Prochaines étapes", level=3)
        for etape in entree["prochaines"]:
            doc.add_paragraph(etape, style="List Bullet")

    # Notes libres
    if entree.get("notes"):
        doc.add_heading("📝 Notes", level=3)
        doc.add_paragraph(entree["notes"])

    # Niveau d'avancement
    niveau = entree.get("niveau", "NORMAL").upper()
    p = doc.add_paragraph()
    run = p.add_run(f"Niveau d'avancement : {niveau}")
    run.bold = True
    run.font.color.rgb = NIVEAU_COLORS.get(niveau, RGBColor(0, 0, 0))

    doc.add_paragraph("")


def git_commit_push(message: str) -> None:
    """Ajoute, commit et push les changements sur GitHub."""
    try:
        subprocess.run(["git", "add", "-A"], check=True)
        subprocess.run(["git", "commit", "-m", message], check=True)
        subprocess.run(["git", "push"], check=True)
        print("✅ GitHub mis à jour.")
    except subprocess.CalledProcessError as e:
        print(f"⚠️ Erreur Git : {e}")


def saisir_liste(prompt: str) -> list[str]:
    """Saisie interactive d'une liste (ligne vide pour terminer)."""
    print(f"\n{prompt} (ligne vide pour terminer)")
    items = []
    while True:
        val = input("  → ").strip()
        if not val:
            break
        items.append(val)
    return items


def main():
    print("=" * 50)
    print("  📓 Mise à jour du Journal de Bord")
    print("=" * 50)

    # Charger ou créer le journal
    if not JOURNAL_PATH.exists():
        doc = init_journal()
    else:
        doc = Document(JOURNAL_PATH)

    # Saisie interactive
    taches = saisir_liste("✅ Tâches accomplies aujourd'hui :")
    if not taches:
        print("Aucune tâche saisie. Abandon.")
        return

    difficultes = saisir_liste("⚠️  Difficultés rencontrées (optionnel) :")
    prochaines  = saisir_liste("🎯 Prochaines étapes (optionnel) :")

    print("\n📝 Notes libres (optionnel, entrée vide pour ignorer) :")
    notes = input("  → ").strip()

    print("\nNiveau d'avancement ? [NORMAL / ATTENTION / BLOQUÉ] (défaut : NORMAL)")
    niveau = input("  → ").strip().upper() or "NORMAL"

    entree = {
        "taches":      taches,
        "difficultes": difficultes,
        "prochaines":  prochaines,
        "notes":       notes,
        "niveau":      niveau,
    }

    # Sauvegarder
    ajouter_entree(doc, entree)
    doc.save(JOURNAL_PATH)
    print(f"\n✅ Journal mis à jour : {JOURNAL_PATH}")

    # Git commit + push
    today_str = date.today().isoformat()
    message = f"journal: mise à jour du {today_str} — {taches[0][:60]}"
    git_commit_push(message)

    print("\n🎉 Tout est à jour !")


if __name__ == "__main__":
    main()
